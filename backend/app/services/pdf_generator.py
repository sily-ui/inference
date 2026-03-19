"""
PDF 报告生成服务
使用 ReportLab 生成美观的 PDF 报告
"""

import os
import re
from typing import List, Dict, Any, Optional
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    PageBreak,
    Table,
    TableStyle,
    ListFlowable,
    ListItem,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus.flowables import HRFlowable

from ..utils.logger import get_logger

logger = get_logger("mirofish.pdf_generator")


class PDFGenerator:
    """PDF 报告生成器"""

    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._register_chinese_fonts()
        self._create_custom_styles()

    def _register_chinese_fonts(self):
        """注册中文字体"""
        font_paths = [
            ("SimHei", "C:/Windows/Fonts/simhei.ttf"),
            ("SimSun", "C:/Windows/Fonts/simsun.ttc"),
            ("Microsoft YaHei", "C:/Windows/Fonts/msyh.ttc"),
            ("Microsoft YaHei Bold", "C:/Windows/Fonts/msyhbd.ttc"),
        ]

        self.chinese_font = None
        self.chinese_font_bold = None

        for font_name, font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    if font_name not in pdfmetrics.getRegisteredFontNames():
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                    if "Bold" in font_name or "Hei" in font_name:
                        self.chinese_font_bold = font_name
                    else:
                        self.chinese_font = font_name
                    logger.debug(f"成功注册字体: {font_name}")
                except Exception as e:
                    logger.warning(f"注册字体失败 {font_name}: {e}")

        if not self.chinese_font:
            self.chinese_font = "Helvetica"
        if not self.chinese_font_bold:
            self.chinese_font_bold = "Helvetica-Bold"

    def _create_custom_styles(self):
        """创建自定义样式"""
        font = self.chinese_font
        font_bold = self.chinese_font_bold

        self.styles.add(
            ParagraphStyle(
                name="ChineseTitle",
                fontName=font_bold,
                fontSize=24,
                leading=32,
                alignment=TA_CENTER,
                spaceAfter=20,
                textColor=HexColor("#1a1a2e"),
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseSubtitle",
                fontName=font,
                fontSize=12,
                leading=18,
                alignment=TA_CENTER,
                spaceAfter=30,
                textColor=HexColor("#666666"),
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseHeading1",
                fontName=font_bold,
                fontSize=18,
                leading=26,
                spaceBefore=20,
                spaceAfter=12,
                textColor=HexColor("#16213e"),
                borderPadding=(0, 0, 5, 0),
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseHeading2",
                fontName=font_bold,
                fontSize=14,
                leading=20,
                spaceBefore=15,
                spaceAfter=8,
                textColor=HexColor("#0f3460"),
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseHeading3",
                fontName=font_bold,
                fontSize=12,
                leading=18,
                spaceBefore=10,
                spaceAfter=6,
                textColor=HexColor("#1a1a2e"),
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseBody",
                fontName=font,
                fontSize=10.5,
                leading=18,
                alignment=TA_JUSTIFY,
                spaceBefore=4,
                spaceAfter=8,
                firstLineIndent=21,
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseBodyNoIndent",
                fontName=font,
                fontSize=10.5,
                leading=18,
                alignment=TA_LEFT,
                spaceBefore=4,
                spaceAfter=8,
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseBullet",
                fontName=font,
                fontSize=10.5,
                leading=16,
                leftIndent=20,
                spaceBefore=2,
                spaceAfter=2,
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseQuote",
                fontName=font,
                fontSize=10,
                leading=16,
                leftIndent=30,
                rightIndent=30,
                spaceBefore=10,
                spaceAfter=10,
                textColor=HexColor("#555555"),
                backColor=HexColor("#f5f5f5"),
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseMeta",
                fontName=font,
                fontSize=9,
                leading=14,
                textColor=HexColor("#888888"),
            )
        )

    def _parse_markdown(self, markdown_text: str) -> List[Dict[str, Any]]:
        """解析 Markdown 文本为结构化内容"""
        elements = []
        lines = markdown_text.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            if line.startswith("# "):
                elements.append({"type": "h1", "content": line[2:].strip()})
            elif line.startswith("## "):
                elements.append({"type": "h2", "content": line[3:].strip()})
            elif line.startswith("### "):
                elements.append({"type": "h3", "content": line[4:].strip()})
            elif line.startswith("#### "):
                elements.append({"type": "h3", "content": line[5:].strip()})
            elif line.startswith("- ") or line.startswith("* "):
                bullet_items = []
                while i < len(lines) and (
                    lines[i].strip().startswith("- ")
                    or lines[i].strip().startswith("* ")
                ):
                    bullet_items.append(lines[i].strip()[2:].strip())
                    i += 1
                elements.append({"type": "bullet", "items": bullet_items})
                continue
            elif line.startswith(("1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "7. ", "8. ", "9. ")):
                numbered_items = []
                while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                    content = re.sub(r"^\d+\.\s*", "", lines[i].strip())
                    numbered_items.append(content)
                    i += 1
                elements.append({"type": "numbered", "items": numbered_items})
                continue
            elif line.startswith("> "):
                quote_lines = []
                while i < len(lines) and lines[i].strip().startswith("> "):
                    quote_lines.append(lines[i].strip()[2:].strip())
                    i += 1
                elements.append({"type": "quote", "content": "\n".join(quote_lines)})
                continue
            elif line.startswith("---") or line.startswith("***"):
                elements.append({"type": "divider"})
            elif line.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                elements.append({"type": "code", "content": "\n".join(code_lines)})
            else:
                paragraph_lines = [line]
                i += 1
                while i < len(lines):
                    next_line = lines[i].strip()
                    if (
                        not next_line
                        or next_line.startswith("#")
                        or next_line.startswith("-")
                        or next_line.startswith("*")
                        or next_line.startswith(">")
                        or next_line.startswith("```")
                        or re.match(r"^\d+\.\s", next_line)
                    ):
                        break
                    paragraph_lines.append(next_line)
                    i += 1
                elements.append({"type": "paragraph", "content": " ".join(paragraph_lines)})
                continue

            i += 1

        return elements

    def _clean_text(self, text: str) -> str:
        """清理文本中的 Markdown 格式标记"""
        text = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", text)
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
        text = re.sub(r"~~(.+?)~~", r"<strike>\1</strike>", text)
        text = re.sub(r"`(.+?)`", r"<font face='Courier' size='9'>\1</font>", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        text = text.replace("<", "&lt;").replace(">", "&gt;")
        text = text.replace("&lt;b&gt;", "<b>").replace("&lt;/b&gt;", "</b>")
        text = text.replace("&lt;i&gt;", "<i>").replace("&lt;/i&gt;", "</i>")
        text = text.replace("&lt;strike&gt;", "<strike>").replace("&lt;/strike&gt;", "</strike>")
        text = text.replace("&lt;font", "<font").replace("&lt;/font&gt;", "</font>")

        return text

    def _add_header_footer(self, canvas, doc):
        """添加页眉页脚"""
        canvas.saveState()

        canvas.setFont(self.chinese_font, 8)
        canvas.setFillColor(HexColor("#888888"))

        canvas.drawString(doc.leftMargin, doc.height + doc.topMargin + 10 * mm, "MiroFish Prediction Report")

        page_num = canvas.getPageNumber()
        canvas.drawRightString(doc.width + doc.leftMargin, 15 * mm, f"- {page_num} -")

        canvas.restoreState()

    def _add_element(self, story: List, elem: Dict[str, Any]):
        """添加元素到故事中"""
        elem_type = elem.get("type")

        if elem_type == "h1":
            story.append(
                Paragraph(self._clean_text(elem["content"]), self.styles["ChineseHeading1"])
            )
            story.append(
                HRFlowable(
                    width="30%",
                    thickness=2,
                    color=HexColor("#16213e"),
                    spaceBefore=2,
                    spaceAfter=10,
                )
            )

        elif elem_type == "h2":
            story.append(
                Paragraph(self._clean_text(elem["content"]), self.styles["ChineseHeading2"])
            )

        elif elem_type == "h3":
            story.append(
                Paragraph(self._clean_text(elem["content"]), self.styles["ChineseHeading3"])
            )

        elif elem_type == "paragraph":
            story.append(
                Paragraph(self._clean_text(elem["content"]), self.styles["ChineseBody"])
            )

        elif elem_type == "bullet":
            for item in elem["items"]:
                story.append(
                    Paragraph(
                        f"• {self._clean_text(item)}", self.styles["ChineseBodyNoIndent"]
                    )
                )

        elif elem_type == "numbered":
            for idx, item in enumerate(elem["items"], 1):
                story.append(
                    Paragraph(
                        f"{idx}. {self._clean_text(item)}", self.styles["ChineseBodyNoIndent"]
                    )
                )

        elif elem_type == "quote":
            story.append(
                Paragraph(
                    f'"{self._clean_text(elem["content"])}"', self.styles["ChineseQuote"]
                )
            )

        elif elem_type == "divider":
            story.append(
                HRFlowable(
                    width="100%",
                    thickness=0.5,
                    color=HexColor("#cccccc"),
                    spaceBefore=10,
                    spaceAfter=10,
                )
            )

        elif elem_type == "code":
            code_text = elem["content"].replace("<", "&lt;").replace(">", "&gt;")
            story.append(
                Paragraph(
                    f"<font face='Courier' size='9'><pre>{code_text}</pre></font>",
                    self.styles["ChineseBodyNoIndent"],
                )
            )

    def generate_pdf(
        self,
        output_path: str,
        title: str,
        subtitle: str = "",
        content: str = "",
        sections: List[Dict[str, Any]] = None,
        metadata: Dict[str, Any] = None,
    ) -> str:
        """
        生成 PDF 报告

        Args:
            output_path: 输出文件路径
            title: 报告标题
            subtitle: 副标题/摘要
            content: Markdown 格式的完整内容
            sections: 分章节内容列表 [{"title": "...", "content": "..."}]
            metadata: 元数据 {"author": "...", "date": "..."}

        Returns:
            生成的 PDF 文件路径
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2 * cm,
        )

        story = []

        story.append(Paragraph(self._clean_text(title), self.styles["ChineseTitle"]))

        if subtitle:
            story.append(
                Paragraph(self._clean_text(subtitle), self.styles["ChineseSubtitle"])
            )

        story.append(Spacer(1, 10 * mm))

        if metadata:
            meta_text = ""
            if metadata.get("author"):
                meta_text += f"作者: {metadata['author']}  "
            if metadata.get("date"):
                meta_text += f"日期: {metadata['date']}  "
            if metadata.get("report_id"):
                meta_text += f"报告ID: {metadata['report_id']}"
            if meta_text:
                story.append(Paragraph(meta_text, self.styles["ChineseMeta"]))
                story.append(Spacer(1, 5 * mm))

        story.append(
            HRFlowable(width="100%", thickness=1, color=HexColor("#e0e0e0"), spaceBefore=5, spaceAfter=15)
        )

        if sections:
            for section in sections:
                section_title = section.get("title", "")
                section_content = section.get("content", "")

                if section_title:
                    story.append(
                        Paragraph(
                            self._clean_text(section_title), self.styles["ChineseHeading1"]
                        )
                    )
                    story.append(
                        HRFlowable(
                            width="30%",
                            thickness=2,
                            color=HexColor("#16213e"),
                            spaceBefore=2,
                            spaceAfter=10,
                        )
                    )

                if section_content:
                    elements = self._parse_markdown(section_content)
                    for elem in elements:
                        self._add_element(story, elem)

                story.append(Spacer(1, 8 * mm))

        elif content:
            elements = self._parse_markdown(content)
            for elem in elements:
                self._add_element(story, elem)

        doc.build(story, onFirstPage=self._add_header_footer, onLaterPages=self._add_header_footer)

        logger.info(f"PDF 报告已生成: {output_path}")
        return output_path


def generate_report_pdf(
    report_id: str,
    title: str,
    summary: str = "",
    markdown_content: str = "",
    sections: List[Dict[str, Any]] = None,
    output_dir: str = None,
) -> str:
    """
    生成报告 PDF 的便捷函数

    Args:
        report_id: 报告ID
        title: 报告标题
        summary: 报告摘要
        markdown_content: Markdown 格式的完整内容
        sections: 分章节内容
        output_dir: 输出目录

    Returns:
        PDF 文件路径
    """
    from ..config import Config

    if output_dir is None:
        output_dir = os.path.join(Config.UPLOAD_FOLDER, "reports", report_id)

    output_path = os.path.join(output_dir, f"{report_id}.pdf")

    generator = PDFGenerator()

    metadata = {
        "report_id": report_id,
        "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }

    generator.generate_pdf(
        output_path=output_path,
        title=title,
        subtitle=summary,
        content=markdown_content,
        sections=sections,
        metadata=metadata,
    )

    return output_path


def generate_report_word(
    report_id: str,
    title: str,
    summary: str = "",
    markdown_content: str = "",
    sections: List[Dict[str, Any]] = None,
    output_dir: str = None,
) -> str:
    """
    生成报告 Word 文档的便捷函数

    Args:
        report_id: 报告ID
        title: 报告标题
        summary: 报告摘要
        markdown_content: Markdown 格式的完整内容
        sections: 分章节内容
        output_dir: 输出目录

    Returns:
        Word 文件路径
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from ..config import Config

    if output_dir is None:
        output_dir = os.path.join(Config.UPLOAD_FOLDER, "reports", report_id)

    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{report_id}.docx")

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(10.5)

    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.name = 'SimHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(26, 26, 46)

    if summary:
        summary_para = doc.add_paragraph(summary)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in summary_para.runs:
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(102, 102, 102)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(f"报告ID: {report_id}  |  生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(136, 136, 136)

    doc.add_paragraph()

    def add_markdown_content(text: str):
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                h = doc.add_heading(line[2:], level=1)
                for run in h.runs:
                    run.font.name = 'SimHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            elif line.startswith('## '):
                h = doc.add_heading(line[3:], level=2)
                for run in h.runs:
                    run.font.name = 'SimHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            elif line.startswith('### '):
                h = doc.add_heading(line[4:], level=3)
                for run in h.runs:
                    run.font.name = 'SimHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                for run in p.runs:
                    run.font.name = 'SimSun'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            elif re.match(r'^\d+\.\s', line):
                content = re.sub(r'^\d+\.\s*', '', line)
                p = doc.add_paragraph(content, style='List Number')
                for run in p.runs:
                    run.font.name = 'SimSun'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            elif line.startswith('> '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(line[2:])
                run.font.name = 'SimSun'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                run.font.italic = True
                run.font.color.rgb = RGBColor(85, 85, 85)
            else:
                clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
                clean_line = re.sub(r'`(.+?)`', r'\1', clean_line)
                p = doc.add_paragraph(clean_line)
                for run in p.runs:
                    run.font.name = 'SimSun'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    if sections:
        for section in sections:
            section_title = section.get("title", "")
            section_content = section.get("content", "")

            if section_title:
                h = doc.add_heading(section_title, level=1)
                for run in h.runs:
                    run.font.name = 'SimHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

            if section_content:
                add_markdown_content(section_content)

    elif markdown_content:
        add_markdown_content(markdown_content)

    doc.save(output_path)
    logger.info(f"Word 报告已生成: {output_path}")
    return output_path
